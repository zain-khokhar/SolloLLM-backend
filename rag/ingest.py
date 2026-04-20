"""
Document Ingestion Pipeline for SoloLLM.

Handles parsing of multiple document formats:
PDF (layout-aware), DOCX, TXT, MD, HTML, CSV, EPUB, code files.
"""

import logging
import os
import re
import csv
import io
from pathlib import Path
from datetime import datetime, timezone
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)


@dataclass
class DocumentSection:
    """A section within a parsed document."""
    title: str
    content: str
    level: int = 0  # heading level (1-6)
    page_number: int | None = None
    section_index: int = 0
    parent_title: str = ""
    metadata: dict = field(default_factory=dict)


@dataclass
class ParsedDocument:
    """Result of parsing a document."""
    filename: str
    file_type: str
    title: str
    content: str  # Full text content
    sections: list[DocumentSection] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)
    page_count: int = 0
    tables: list[dict] = field(default_factory=list)
    errors: list[str] = field(default_factory=list)

    @property
    def total_chars(self) -> int:
        return len(self.content)


# ── File type detection ─────────────────────────────────────

SUPPORTED_EXTENSIONS = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".doc": "doc",
    ".txt": "text",
    ".md": "markdown",
    ".markdown": "markdown",
    ".html": "html",
    ".htm": "html",
    ".csv": "csv",
    ".tsv": "csv",
    ".py": "code",
    ".js": "code",
    ".ts": "code",
    ".jsx": "code",
    ".tsx": "code",
    ".java": "code",
    ".c": "code",
    ".cpp": "code",
    ".h": "code",
    ".rs": "code",
    ".go": "code",
    ".rb": "code",
    ".php": "code",
    ".json": "code",
    ".yaml": "code",
    ".yml": "code",
    ".xml": "code",
    ".sql": "code",
    ".sh": "code",
    ".bat": "code",
    ".ps1": "code",
}


def detect_file_type(filename: str) -> str | None:
    """Detect document type from file extension."""
    ext = Path(filename).suffix.lower()
    return SUPPORTED_EXTENSIONS.get(ext)


# ── Parsers ─────────────────────────────────────────────────

def parse_text(content: str, filename: str) -> ParsedDocument:
    """Parse plain text files."""
    return ParsedDocument(
        filename=filename,
        file_type="text",
        title=Path(filename).stem,
        content=content,
        sections=[DocumentSection(title="Content", content=content, level=0)],
    )


def parse_markdown(content: str, filename: str) -> ParsedDocument:
    """Parse markdown files with section detection."""
    sections = []
    current_section = None
    section_index = 0

    lines = content.split("\n")
    buffer = []
    title = Path(filename).stem

    for line in lines:
        heading_match = re.match(r'^(#{1,6})\s+(.+)', line)
        if heading_match:
            # Save previous section
            if current_section and buffer:
                current_section.content = "\n".join(buffer).strip()
                sections.append(current_section)
                buffer = []

            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()

            if level == 1 and not sections:
                title = heading_text

            current_section = DocumentSection(
                title=heading_text,
                content="",
                level=level,
                section_index=section_index,
            )
            section_index += 1
        else:
            buffer.append(line)

    # Save last section
    if current_section and buffer:
        current_section.content = "\n".join(buffer).strip()
        sections.append(current_section)
    elif buffer and not sections:
        sections.append(DocumentSection(
            title="Content",
            content="\n".join(buffer).strip(),
            level=0,
        ))

    return ParsedDocument(
        filename=filename,
        file_type="markdown",
        title=title,
        content=content,
        sections=sections,
    )


def parse_html(content: str, filename: str) -> ParsedDocument:
    """Parse HTML files — strip tags, extract text."""
    # Simple HTML tag stripping without external dependency
    text = re.sub(r'<script[^>]*>.*?</script>', '', content, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<[^>]+>', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()

    # Try to extract title
    title_match = re.search(r'<title[^>]*>(.*?)</title>', content, re.IGNORECASE | re.DOTALL)
    title = title_match.group(1).strip() if title_match else Path(filename).stem

    return ParsedDocument(
        filename=filename,
        file_type="html",
        title=title,
        content=text,
        sections=[DocumentSection(title="Content", content=text, level=0)],
    )


def parse_csv(content: str, filename: str) -> ParsedDocument:
    """Parse CSV files into markdown tables."""
    reader = csv.reader(io.StringIO(content))
    rows = list(reader)

    if not rows:
        return ParsedDocument(
            filename=filename,
            file_type="csv",
            title=Path(filename).stem,
            content="(Empty CSV file)",
            sections=[],
        )

    # Build markdown table
    headers = rows[0]
    md_lines = []
    md_lines.append("| " + " | ".join(headers) + " |")
    md_lines.append("| " + " | ".join(["---"] * len(headers)) + " |")

    for row in rows[1:]:
        # Pad row to match header length
        padded = row + [""] * (len(headers) - len(row))
        md_lines.append("| " + " | ".join(padded[:len(headers)]) + " |")

    md_content = "\n".join(md_lines)

    return ParsedDocument(
        filename=filename,
        file_type="csv",
        title=Path(filename).stem,
        content=md_content,
        sections=[DocumentSection(title="Data", content=md_content, level=0)],
        tables=[{"headers": headers, "rows": rows[1:], "source": filename}],
        metadata={"row_count": len(rows) - 1, "column_count": len(headers)},
    )


def parse_code(content: str, filename: str) -> ParsedDocument:
    """Parse source code files with function/class detection."""
    ext = Path(filename).suffix.lower()
    lang = {
        ".py": "python", ".js": "javascript", ".ts": "typescript",
        ".jsx": "jsx", ".tsx": "tsx", ".java": "java",
        ".c": "c", ".cpp": "cpp", ".h": "c", ".rs": "rust",
        ".go": "go", ".rb": "ruby", ".php": "php",
    }.get(ext, "text")

    # Wrap in code block for context
    wrapped = f"```{lang}\n{content}\n```"

    return ParsedDocument(
        filename=filename,
        file_type="code",
        title=Path(filename).stem,
        content=content,
        sections=[DocumentSection(
            title=filename,
            content=wrapped,
            level=0,
            metadata={"language": lang},
        )],
        metadata={"language": lang, "line_count": content.count("\n") + 1},
    )


def parse_pdf(file_path: str) -> ParsedDocument:
    """
    Parse PDF files with layout awareness.
    Uses PyMuPDF (fitz) for text extraction with structure detection.
    """
    filename = os.path.basename(file_path)

    try:
        import fitz  # PyMuPDF
    except ImportError:
        return ParsedDocument(
            filename=filename,
            file_type="pdf",
            title=Path(filename).stem,
            content="",
            errors=["PyMuPDF (fitz) is not installed. Run: pip install PyMuPDF"],
        )

    try:
        sections = []
        all_text = []
        tables = []
        page_count = 0

        with fitz.open(file_path) as doc:
            page_count = len(doc)

            metadata = {}
            if doc.metadata:
                metadata = {
                    "author": doc.metadata.get("author", ""),
                    "title": doc.metadata.get("title", ""),
                    "subject": doc.metadata.get("subject", ""),
                    "creator": doc.metadata.get("creator", ""),
                    "page_count": page_count,
                }

            title = metadata.get("title") or Path(filename).stem

            pdf_plumber_doc = None
            try:
                import pdfplumber
                pdf_plumber_doc = pdfplumber.open(file_path)
            except ImportError:
                pdf_plumber_doc = None
            except Exception:
                pdf_plumber_doc = None

            for page_num, page in enumerate(doc, 1):
                page_text = page.get_text("text")
                if page_text.strip():
                    all_text.append(page_text)

                    # Detect sections by looking for bold/large text patterns
                    blocks = page.get_text("dict")["blocks"]
                    for block in blocks:
                        if block.get("type") == 0:  # text block
                            for line in block.get("lines", []):
                                line_text = ""
                                max_size = 0
                                is_bold = False
                                for span in line.get("spans", []):
                                    line_text += span.get("text", "")
                                    size = span.get("size", 12)
                                    if size > max_size:
                                        max_size = size
                                    if "bold" in span.get("font", "").lower():
                                        is_bold = True

                                line_text = line_text.strip()
                                # Detect headings (large or bold text, short lines)
                                if line_text and (max_size > 14 or is_bold) and len(line_text) < 200:
                                    level = 1 if max_size > 20 else (2 if max_size > 16 else 3)
                                    sections.append(DocumentSection(
                                        title=line_text,
                                        content="",
                                        level=level,
                                        page_number=page_num,
                                        section_index=len(sections),
                                    ))

                    # Always add a stable page-based section with content.
                    # Heading-only sections can be sparse and cause poor downstream chunking.
                    sections.append(DocumentSection(
                        title=f"Page {page_num}",
                        content=page_text.strip(),
                        level=0,
                        page_number=page_num,
                        section_index=len(sections),
                    ))

                # Try to extract tables using pdfplumber if available
                if pdf_plumber_doc is not None:
                    try:
                        if page_num <= len(pdf_plumber_doc.pages):
                            p = pdf_plumber_doc.pages[page_num - 1]
                            page_tables = p.extract_tables()
                            for table in page_tables:
                                if table and len(table) > 1:
                                    headers = [str(c or "") for c in table[0]]
                                    rows = [[str(c or "") for c in row] for row in table[1:]]
                                    tables.append({
                                        "headers": headers,
                                        "rows": rows,
                                        "page": page_num,
                                        "source": filename,
                                    })
                    except Exception:
                        pass

            if pdf_plumber_doc is not None:
                try:
                    pdf_plumber_doc.close()
                except Exception:
                    pass

        full_text = "\n\n".join(all_text)

        # Keep only sections that have real content to avoid empty/duplicate structural noise.
        sections = [s for s in sections if s.content and s.content.strip()]

        return ParsedDocument(
            filename=filename,
            file_type="pdf",
            title=title,
            content=full_text,
            sections=sections,
            metadata=metadata,
            page_count=page_count or metadata.get("page_count", 0),
            tables=tables,
        )

    except Exception as e:
        logger.error(f"Failed to parse PDF {filename}: {e}")
        return ParsedDocument(
            filename=filename,
            file_type="pdf",
            title=Path(filename).stem,
            content="",
            errors=[f"PDF parsing failed: {str(e)}"],
        )


def parse_docx(file_path: str) -> ParsedDocument:
    """Parse DOCX files with structure detection."""
    filename = os.path.basename(file_path)

    try:
        from docx import Document as DocxDocument
    except ImportError:
        return ParsedDocument(
            filename=filename,
            file_type="docx",
            title=Path(filename).stem,
            content="",
            errors=["python-docx is not installed. Run: pip install python-docx"],
        )

    try:
        doc = DocxDocument(file_path)
        sections = []
        all_text = []
        tables = []

        title = Path(filename).stem
        if doc.core_properties.title:
            title = doc.core_properties.title

        # Parse paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            all_text.append(text)

            # Detect headings by style
            style_name = (para.style.name or "").lower()
            if "heading" in style_name:
                try:
                    level = int(style_name.replace("heading", "").strip())
                except ValueError:
                    level = 1
                sections.append(DocumentSection(
                    title=text,
                    content="",
                    level=level,
                    section_index=len(sections),
                ))
            elif sections:
                # Append to current section
                if sections[-1].content:
                    sections[-1].content += "\n" + text
                else:
                    sections[-1].content = text

        # Parse tables
        for table in doc.tables:
            rows_data = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows_data.append(cells)
            if rows_data and len(rows_data) > 1:
                tables.append({
                    "headers": rows_data[0],
                    "rows": rows_data[1:],
                    "source": filename,
                })

        full_text = "\n\n".join(all_text)

        # If no headings found, create a single section
        if not sections:
            sections = [DocumentSection(
                title="Content",
                content=full_text,
                level=0,
            )]

        return ParsedDocument(
            filename=filename,
            file_type="docx",
            title=title,
            content=full_text,
            sections=sections,
            tables=tables,
            metadata={
                "author": doc.core_properties.author or "",
            },
        )

    except Exception as e:
        logger.error(f"Failed to parse DOCX {filename}: {e}")
        return ParsedDocument(
            filename=filename,
            file_type="docx",
            title=Path(filename).stem,
            content="",
            errors=[f"DOCX parsing failed: {str(e)}"],
        )


# ── Main ingestion function ─────────────────────────────────

async def ingest_file(file_path: str) -> ParsedDocument:
    """
    Ingest a document file and return parsed content.
    Automatically detects file type and uses the appropriate parser.
    """
    filename = os.path.basename(file_path)
    file_type = detect_file_type(filename)

    if not file_type:
        return ParsedDocument(
            filename=filename,
            file_type="unknown",
            title=Path(filename).stem,
            content="",
            errors=[f"Unsupported file type: {Path(filename).suffix}"],
        )

    logger.info(f"Ingesting document: {filename} (type: {file_type})")

    if file_type == "pdf":
        return parse_pdf(file_path)
    elif file_type == "docx":
        return parse_docx(file_path)
    elif file_type in ("text", "markdown", "html", "csv", "code"):
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                content = f.read()
        except Exception as e:
            return ParsedDocument(
                filename=filename,
                file_type=file_type,
                title=Path(filename).stem,
                content="",
                errors=[f"Failed to read file: {str(e)}"],
            )

        if file_type == "markdown":
            return parse_markdown(content, filename)
        elif file_type == "html":
            return parse_html(content, filename)
        elif file_type == "csv":
            return parse_csv(content, filename)
        elif file_type == "code":
            return parse_code(content, filename)
        else:
            return parse_text(content, filename)
    else:
        return ParsedDocument(
            filename=filename,
            file_type=file_type,
            title=Path(filename).stem,
            content="",
            errors=[f"Parser not yet implemented for: {file_type}"],
        )
